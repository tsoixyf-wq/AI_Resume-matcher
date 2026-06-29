"""
Multi-format document loader for resumes.
Supports PDF, DOCX, TXT, MD files.

CPU-bound parsing methods are offloaded to a thread pool via ``asyncio.to_thread``
to avoid blocking the FastAPI event loop on large documents.
"""

import asyncio
import os
from pathlib import Path

import structlog

logger = structlog.get_logger(__name__)


class DocumentLoader:
    """Load and extract text from various document formats."""

    @staticmethod
    async def load(file_path: str) -> str:
        """Load a document and extract its full text.

        Args:
            file_path: Path to the document file.

        Returns:
            Extracted plain text content.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = Path(file_path).suffix.lower()

        loaders = {
            ".pdf": DocumentLoader._load_pdf,
            ".docx": DocumentLoader._load_docx,
            ".doc": DocumentLoader._load_docx,
            ".txt": DocumentLoader._load_txt,
            ".md": DocumentLoader._load_txt,
        }

        loader = loaders.get(ext)
        if loader is None:
            raise ValueError(f"Unsupported file format: {ext}")

        try:
            text = await asyncio.to_thread(loader, file_path)
            logger.info("Document loaded", file=file_path, ext=ext, length=len(text))
            return text.strip()
        except Exception as e:
            logger.error("Failed to load document", file=file_path, error=str(e))
            raise

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        """Extract text from PDF using PyMuPDF (fitz)."""
        import fitz  # PyMuPDF

        doc = fitz.open(file_path)
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        return "\n\n".join(text_parts)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        """Extract text from DOCX file."""
        from docx import Document

        doc = Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n".join(text_parts)

    @staticmethod
    def _load_txt(file_path: str) -> str:
        """Extract text from plain text file."""
        with open(file_path, encoding="utf-8", errors="replace") as f:
            return f.read()
